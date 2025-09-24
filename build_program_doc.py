from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUT_FILE = "CartPole_DQN_DDQN_Program_Doc.docx"

# ---------------- helpers ----------------
def ensure_styles(doc):
    styles = doc.styles
    if "CaptionCenter" not in styles:
        s = styles.add_style("CaptionCenter", WD_STYLE_TYPE.PARAGRAPH)
        s.font.italic = True
        s.font.size = Pt(10)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_picture_if_exists(doc, path, width_in=6.5, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        if caption:
            cap = doc.add_paragraph(caption, style="CaptionCenter")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_top_level_pictures(doc, folder, width_in=6.5, exclude=()):
    """
    Insert all .png files from the top-level 'folder', excluding any
    filenames listed in 'exclude'. No recursion.
    """
    exclude_set = {f.lower() for f in exclude}
    pngs = [
        f for f in sorted(os.listdir(folder))
        if f.lower().endswith(".png") and f.lower() not in exclude_set
    ]
    for fname in pngs:
        path = os.path.join(folder, fname)
        doc.add_picture(path, width=Inches(width_in))
        cap = doc.add_paragraph(os.path.splitext(fname)[0], style="CaptionCenter")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --------------- document build ---------------
doc = Document()
ensure_styles(doc)

# Title & author
add_heading(doc, "CartPole: Deep Q-Learning vs Double DQN on CPU and GPU", level=0)
add_para(doc, "Author: Antwain M. Sparks and Vojislav Stojkovic, Morgan State University, Computer Science Department, Baltimore, MD 21251")

# Abstract
add_heading(doc, "Abstract", level=1)
add_para(
    doc,
    "Reinforcement learning (RL) is a branch of machine learning in which an agent learns decision-"
    "making through interaction with an environment. Q-Learning (QL) is a foundational RL algorithm "
    "that assigns values to state–action pairs without requiring a model of the environment. "
    "Deep Q-Learning (DQL) extends QL by combining it with deep neural networks, enabling agents to "
    "handle larger and more complex state spaces."
)
add_para(
    doc,
    "The CartPole environment serves as a classic benchmark for evaluating RL methods, where an agent "
    "must balance an inverted pendulum by applying discrete forces to a cart. Despite its apparent "
    "simplicity, CartPole provides a meaningful platform for comparing algorithm performance."
)
add_para(
    doc,
    "We implement and evaluate DQL under three configurations: a CPU-based baseline, a GPU-accelerated "
    "DQN, and a GPU-accelerated Double DQN (DDQN). While GPU acceleration speeds up training, "
    "algorithmic improvements such as DDQN are necessary for stability and convergence. The combined "
    "approach (GPU + DDQN) consistently solves CartPole at the maximum reward threshold."
)

# Methodology
add_heading(doc, "Methodology", level=1)
add_para(doc, "• CartPole environment setup")
add_para(doc, "• DQN vs DDQN algorithms (target network decouples action selection and evaluation)")
add_para(doc, "• CPU vs GPU execution differences")
add_para(doc, "• Training setup: 800 episodes, epsilon decay to 0.01, replay buffer=100k, batch=64, lr=5e-4")

# Results – Table (plain text for the program doc)
add_heading(doc, "Results (Summary)", level=1)
add_para(doc, "Model\t\t\tDevice\tEpisodes to Solve\tFinal Avg Reward\tTraining Time (mins)")
add_para(doc, "DQN (CPU baseline)\tCPU\t>600 (unstable)\t\t~183\t\t\t~8.5")
add_para(doc, "DQN (GPU)\t\tGPU\t800 (not solved)\t\t288.6\t\t\t~7")
add_para(doc, "DDQN (GPU)\t\tGPU\t~480 (solved)\t\t500.0\t\t\t~8.5")

# Charts — show first
add_heading(doc, "Charts", level=1)
# Put charts first in this explicit order
charts_first = [
    ("comparison.png", "Training rewards (moving avg) — CPU vs GPU DQN vs GPU DDQN"),
    ("eval_comparison.png", "Evaluation comparison: average return (CPU DQN vs GPU DQN vs GPU DDQN)")
]
for fname, caption in charts_first:
    add_picture_if_exists(doc, fname, width_in=6.5, caption=caption)

# Figures — then add remaining PNGs from top level, excluding the charts so we don't duplicate
add_heading(doc, "Figures (Screenshots & Extras)", level=1)
add_top_level_pictures(
    doc,
    folder=".",
    width_in=6.5,
    exclude=[name for name, _ in charts_first]
)

# Technical Specs
add_heading(doc, "Technical Specs", level=1)
specs = [
    "CPU: Intel i7-12700H (20 cores)",
    "RAM: 16 GB",
    "GPU: NVIDIA GeForce RTX 3050 Ti Laptop GPU (4 GB VRAM)",
    "CUDA: 12.4",
    "Python: 3.11.9, PyTorch 2.6.0+cu124",
]
for s in specs:
    add_para(doc, f"• {s}")

# Key Takeaways
add_heading(doc, "Key Takeaways", level=1)
add_para(doc, "• CPU DQN: slow, unstable, fails to solve.")
add_para(doc, "• GPU DQN: faster but underperforms vs DDQN.")
add_para(doc, "• GPU DDQN: consistently solves CartPole (avg return 500).")
add_para(doc, "• Lesson: Algorithmic improvements (DDQN) + GPU acceleration deliver the best results.")

# Conclusion
add_heading(doc, "Conclusion", level=1)
add_para(
    doc,
    "GPU acceleration helps, but algorithmic refinement (DDQN vs DQN) is essential to reliably solve "
    "CartPole. The combination of DDQN with GPU provided the strongest performance."
)

# Save
doc.save(OUT_FILE)
print(f"Saved program doc to: {OUT_FILE}")
